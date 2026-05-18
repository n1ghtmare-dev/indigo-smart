"""Insert figures into the diploma docx before each 'Рисунок N – ...' caption."""
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

DOC = r"c:\www-Oleg\IndigoSmart\docs\diploma\Дипломная_Работа_ИСиП_2_22_Боев_Олег.docx"
BASE = r"c:\www-Oleg\IndigoSmart\docs\figures"

# Map figure number -> image path
FIGURES = {
    1: os.path.join(BASE, "fig1.png"),
    2: os.path.join(BASE, "fig2.png"),
    3: os.path.join(BASE, "fig3.png"),
    4: os.path.join(BASE, "fig4.png"),
    5: os.path.join(BASE, "final-dashboard.png"),
    6: os.path.join(BASE, "final-analytics.png"),
}

# Different widths: diagrams wider, UI screenshots can stay narrower (taller)
WIDTHS = {
    1: Inches(6.5),
    2: Inches(6.5),
    3: Inches(6.5),
    4: Inches(6.5),
    5: Inches(5.5),  # tall mobile-like screenshot
    6: Inches(5.5),
}

# Verify all images exist
for fnum, path in FIGURES.items():
    if not os.path.exists(path):
        raise FileNotFoundError(f"Missing image for figure {fnum}: {path}")

doc = Document(DOC)

# Pattern to match "Рисунок N – ..." captions
caption_pat = re.compile(r"^Рисунок\s+(\d+)\s*[–-]")

inserted_count = 0
for i, par in enumerate(doc.paragraphs):
    m = caption_pat.match(par.text.strip())
    if not m:
        continue
    fnum = int(m.group(1))
    if fnum not in FIGURES:
        continue

    # Take previous paragraph as image holder; if it has content, insert a new one
    prev_idx = i - 1
    if prev_idx < 0:
        continue
    holder = doc.paragraphs[prev_idx]
    # If holder already has text or runs, insert a fresh empty paragraph before the caption
    if holder.text.strip() or holder.runs:
        holder = par.insert_paragraph_before("", style=par.style)

    # Clear any leftover runs from holder
    for r in list(holder.runs):
        r._r.getparent().remove(r._r)
    holder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = holder.add_run()
    run.add_picture(FIGURES[fnum], width=WIDTHS[fnum])
    inserted_count += 1
    print(f"  Inserted figure {fnum} before paragraph {i}: '{par.text.strip()[:60]}'")

try:
    doc.save(DOC)
    print(f"\nDone. Inserted {inserted_count} figures into {DOC}")
except PermissionError:
    alt = DOC.replace(".docx", ".WITH_FIGURES.docx")
    doc.save(alt)
    print(f"\n[Original file is open in Word, saved as a sibling instead]")
    print(f"Saved: {alt}")
    print("Close Word and replace the original with this file, OR re-run after closing.")

