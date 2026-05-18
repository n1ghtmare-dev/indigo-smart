"""Analyze diploma docx structure, styles, fonts, and headings."""
from docx import Document
from collections import Counter
import sys
import io

DOC = r"c:\www-Oleg\IndigoSmart\docs\diploma\Дипломная_Работа_ИСиП_2_22_Боев_Олег.docx"
OUT = r"c:\www-Oleg\IndigoSmart\docs\diploma\diploma_analysis.txt"

doc = Document(DOC)
out = io.StringIO()

out.write("=== SECTIONS ===\n")
for i, section in enumerate(doc.sections):
    out.write(f"Section {i}: page_size={section.page_width}x{section.page_height}, "
              f"margins L/R/T/B={section.left_margin}/{section.right_margin}/{section.top_margin}/{section.bottom_margin}\n")

out.write("\n=== STYLE INVENTORY (paragraph styles used) ===\n")
style_counter = Counter()
for p in doc.paragraphs:
    style_counter[p.style.name] += 1
for s, c in style_counter.most_common():
    out.write(f"  {s}: {c}\n")

out.write(f"\n=== TABLES === total: {len(doc.tables)}\n")
for ti, t in enumerate(doc.tables):
    out.write(f"  Table {ti}: rows={len(t.rows)}, cols={len(t.columns)}\n")
    for ri, row in enumerate(t.rows[:3]):
        cells = [c.text.strip()[:50] for c in row.cells]
        out.write(f"    row {ri}: {cells}\n")

out.write("\n=== FULL TEXT OUTLINE (with style + length) ===\n")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    style = p.style.name
    font_info = ""
    if p.runs:
        r = p.runs[0]
        sz = r.font.size.pt if r.font.size else None
        font_info = f"[font={r.font.name or '-'}, sz={sz}, b={r.bold}, i={r.italic}]"
    snippet = text[:200].replace('\n', ' ')
    out.write(f"{i:4}  {style:20}  {font_info:50}  {snippet}\n")

with open(OUT, "w", encoding="utf-8") as f:
    f.write(out.getvalue())

print(f"Written {len(out.getvalue())} chars to {OUT}")
