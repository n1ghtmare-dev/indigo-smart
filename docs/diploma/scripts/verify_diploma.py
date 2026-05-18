"""Verify the rewritten diploma."""
from docx import Document
import io

DOC = r"c:\www-Oleg\IndigoSmart\docs\diploma\Дипломная_Работа_ИСиП_2_22_Боев_Олег.docx"
OUT = r"c:\www-Oleg\IndigoSmart\docs\diploma\diploma_verify.txt"

doc = Document(DOC)
out = io.StringIO()

out.write(f"=== TABLES === total: {len(doc.tables)}\n")
for ti, t in enumerate(doc.tables):
    if ti == 2:  # Тема
        out.write(f"\nTable 2 (Тема):\n")
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text.strip()
                if txt:
                    out.write(f"  [{ri},{ci}] {txt[:200]}\n")

out.write(f"\n=== TOTAL PARAGRAPHS: {len(doc.paragraphs)} ===\n")

out.write("\n=== HEADINGS ONLY ===\n")
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith("Heading") or (p.text.strip() in ("ВВЕДЕНИЕ", "ЗАКЛЮЧЕНИЕ", "СПИСОК ИСПОЛЬЗУЕМЫХ ИСТОЧНИКОВ")):
        out.write(f"  {i:4}  [{p.style.name}]  {p.text.strip()[:120]}\n")

out.write("\n=== FIRST 5 BODY PARAGRAPHS (after ВВЕДЕНИЕ) ===\n")
found_intro = False
count = 0
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == "ВВЕДЕНИЕ":
        found_intro = True
        continue
    if found_intro and p.text.strip():
        out.write(f"  {i:4}  [{p.style.name}]  {p.text.strip()[:200]}\n")
        count += 1
        if count >= 5:
            break

out.write("\n=== CAPTIONS ===\n")
for i, p in enumerate(doc.paragraphs):
    if p.style.name == "Caption":
        out.write(f"  {i:4}  {p.text.strip()[:120]}\n")

out.write("\n=== LAST 10 PARAGRAPHS ===\n")
for i, p in enumerate(doc.paragraphs[-10:]):
    out.write(f"  [{p.style.name}]  {p.text.strip()[:200]}\n")

with open(OUT, "w", encoding="utf-8") as f:
    f.write(out.getvalue())

print(f"Written verification to {OUT}")
